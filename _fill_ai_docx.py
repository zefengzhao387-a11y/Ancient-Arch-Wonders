# -*- coding: utf-8 -*-
"""Fill competition AI disclosure Word template. Run: python _fill_ai_docx.py"""
from docx import Document

PATH_IN = r"D:\Desktop\2-AI工具使用说明（选用模板）（2026年版）.docx"
PATH = r"D:\Desktop\2-AI工具使用说明（2026年版）-已填写.docx"

FILL_INSTRUCTIONS = """填写说明：
1.本文档适用于所有涉及AI工具使用的参赛作品，表格可另加行。
2.参赛作品的作者，需根据实际的使用情况简明扼要地列出本作品所使用的全部AI工具的名称、版本、访问方式、使用时间、使用环节与目的、关键提示词、AI回复的关键内容、采纳和人工修改情况等。
3.AI回复的关键内容佐证材料，需作为本文档的附录2给出，包括但不限于：（1）关键操作截图（含时间戳，需清晰可辨）；（2）交互录屏视频（时长≤5分钟，需标注使用节点，文档为MP4格式，命名格式：AI_使用序号_作品编号.mp4）；（3）代码注释中标明AI辅助部分（如：// AI辅助生成：DeepSeek-R1-0528, 2025-11-03）
4.提交时，需将本文档的PDF格式文件，以及其他佐证材料（如交互录屏视频），一并上传到作品文件夹的“03设计与开发文档”子文件夹中。
5.本文档内容是正式参赛内容组成部分，需真实填写。如不属实，将导致奖项等级降低甚至终止本作品参加比赛。"""


def delete_table_row(table, row_index: int) -> None:
    tbl = table._tbl
    tr = table.rows[row_index]._tr
    tbl.remove(tr)


def main() -> None:
    d = Document(PATH_IN)

    # 作品编号、名称
    for p in d.paragraphs:
        t = p.text.strip()
        if t.startswith("作品编号：") or "作品编号" in t and "作品名称" in t:
            p.clear()
            p.add_run(
                "作品编号：2026032574    作品名称：古建奇境（Ancient Arch Wonders）"
            )
            break

    # 在「作品编号」段落后插入填写说明（若第3段为空则插入到段首前）
    # 找到作品编号段落索引
    idx = None
    for i, p in enumerate(d.paragraphs):
        if "2026032574" in p.text and "古建奇境" in p.text:
            idx = i
            break
    if idx is not None and idx + 1 < len(d.paragraphs):
        anchor = d.paragraphs[idx + 1]
        if not anchor.text.strip().startswith("填写说明"):
            parts = [ln.strip() for ln in FILL_INSTRUCTIONS.split("\n") if ln.strip()]
            # 单段内换行，避免多段 insert 顺序错乱
            block = "\n".join(parts)
            anchor.insert_paragraph_before(block)

    # 附录1：示例文件夹编号与作品一致
    rep = [
        ("2026012345", "2026032574"),
    ]
    for p in d.paragraphs:
        for a, b in rep:
            if a in p.text:
                p.text = p.text.replace(a, b)

    # 附录2：在「序号N的佐证材料：」下增加一行提示（若尚无子说明）
    # 重新遍历段落，在特定行后插入
    hints = {
        "序号1的佐证材料：": "（佐证建议：含时间戳的 Cursor 对话截图；GameUISfxHub/PersistentGameBGM 相关代码变更说明；录屏命名示例：AI_01_2026032574.mp4）",
        "序号2的佐证材料：": "（佐证建议：CanvasScaler/GlobalCanvasAdaptation 相关讨论截图；脚本中含 // AI辅助 注释的片段截图）",
        "序号3的佐证材料：": "（佐证建议：视频编码问题讨论截图；ffmpeg 转码记录或资源替换说明）",
        "序号4的佐证材料：": "（佐证建议：提交目录 readme 相关对话截图；各 readme.txt 打包目录截图）",
        "序号5的佐证材料：": "（佐证建议：16:9 与 UI 适配讨论截图；可选）",
        "序号6的佐证材料：": "（本序号未使用 AI 可注明「无」；或删除此行提示）",
        "序号7的佐证材料：": "（本序号未使用可注明「无」）",
        "序号8的佐证材料：": "（本序号未使用可注明「无」）",
        "序号9的佐证材料：": "（本序号未使用可注明「无」）",
    }
    inserts = []
    for i, p in enumerate(d.paragraphs):
        text = p.text.strip()
        if text not in hints:
            continue
        if i + 1 < len(d.paragraphs) and d.paragraphs[i + 1].text.strip().startswith(
            "（佐证建议"
        ):
            continue
        inserts.append((i + 1, hints[text]))
    for insert_idx, hint in reversed(inserts):
        d.paragraphs[insert_idx].insert_paragraph_before(hint)

    # 表格：删除空行（第10行，全空）与「示例」行（原第11行，删除后索引会变）
    tbl = d.tables[0]
    # 找到示例行
    example_row = None
    empty_row = None
    for ri, row in enumerate(tbl.rows):
        c0 = row.cells[0].text.strip()
        if c0 == "示例":
            example_row = ri
        if ri > 0 and not any(row.cells[j].text.strip() for j in range(7)):
            # 全空且不是表头
            if c0 == "":
                empty_row = ri
    # 先删示例再删空，避免索引错乱：从大到小删
    to_del = sorted([x for x in (example_row, empty_row) if x is not None], reverse=True)
    for ri in to_del:
        delete_table_row(tbl, ri)

    d.save(PATH)
    print("OK saved:", PATH)
    print("（若需覆盖原模板，请先关闭 Word 后改 PATH 为 PATH_IN 再运行）")


if __name__ == "__main__":
    main()
